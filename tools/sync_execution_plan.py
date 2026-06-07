"""Sync docs/EXECUTION_PLAN.md → docs/EXECUTION_PLAN.docx (+ optional PDF).

Run at start of every session per operator directive 2026-06-06.

Operator-readable Word doc that mirrors the markdown source.  MD remains
source-of-truth; DOCX is regenerated each session for human reading.

Usage:
    py -3.13 tools\\sync_execution_plan.py            # MD -> DOCX
    py -3.13 tools\\sync_execution_plan.py --pdf      # also generate PDF

Requirements:
    python-docx (for DOCX)
    reportlab   (for PDF, optional)
"""

from __future__ import annotations

import argparse
import datetime as _dt
import re
import sys
from pathlib import Path

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
MD_PATH = DOCS / "EXECUTION_PLAN.md"
DOCX_PATH = DOCS / "EXECUTION_PLAN.docx"
PDF_PATH = DOCS / "EXECUTION_PLAN.pdf"

# ---------------------------------------------------------------------------
# Markdown -> DOCX
# ---------------------------------------------------------------------------

def md_to_docx(md_text: str, out_path: Path) -> None:
    """Render a subset of Markdown into a Word document.

    Supported:
      - # H1, ## H2, ### H3
      - **bold** inline runs
      - `code` inline runs
      - Bullet lists (- ...)
      - Numbered lists (1. ...)
      - Checkbox lists (- [ ] / - [x])
      - Pipe tables
      - Fenced code blocks (``` ... ```)
      - Horizontal rules (---)
      - Blockquotes (> ...)
      - Paragraphs

    Anything not in this list renders as a plain paragraph with inline
    formatting preserved.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor, Inches

    doc = Document()

    # ----- Style setup -----
    # Body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Heading sizes
    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        s = doc.styles[name]
        s.font.size = Pt(size)
        s.font.bold = True

    # Page margins tighter for readability
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    # ----- Header / footer with generation timestamp -----
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.text = (
        f"Lyra-cpp Execution Plan — auto-generated "
        f"{_dt.datetime.now().strftime('%Y-%m-%d %H:%M')} "
        f"(source: docs/EXECUTION_PLAN.md)"
    )
    header_p.runs[0].font.size = Pt(8)
    header_p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ----- Parse + render -----
    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    def add_inline_runs(paragraph, text: str) -> None:
        """Render **bold**, `code`, and emoji inline-runs into a paragraph."""
        # Tokenize: split on ** and ` while preserving them
        pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
        parts = pattern.split(text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            else:
                paragraph.add_run(part)

    def shade_cell(cell, hex_color: str) -> None:
        """Apply background shading to a docx table cell."""
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # --- Horizontal rule ---
        if stripped == "---":
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "808080")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # --- Fenced code block ---
        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing ```
            p = doc.add_paragraph()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F4F4F4")
            p._p.get_or_add_pPr().append(shd)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # --- Pipe table ---
        if "|" in stripped and i + 1 < n and re.match(r"^\s*\|?\s*[-:|\s]+\|", lines[i + 1]):
            # Collect rows
            rows: list[list[str]] = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            # rows[0] = header, rows[1] = separator, rows[2..] = body
            if len(rows) >= 2:
                header = rows[0]
                body = rows[2:]
                table = doc.add_table(rows=1 + len(body), cols=len(header))
                table.style = "Light Grid Accent 1"
                # Header row
                hdr_cells = table.rows[0].cells
                for j, h in enumerate(header):
                    if j < len(hdr_cells):
                        cell = hdr_cells[j]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        run = p.add_run(h)
                        run.bold = True
                        run.font.size = Pt(9)
                        shade_cell(cell, "D9E1F2")
                # Body rows
                for r_idx, row in enumerate(body, start=1):
                    body_cells = table.rows[r_idx].cells
                    for j, cell_text in enumerate(row):
                        if j < len(body_cells):
                            cell = body_cells[j]
                            cell.text = ""
                            p = cell.paragraphs[0]
                            add_inline_runs(p, cell_text)
                            for run in p.runs:
                                run.font.size = Pt(9)
                doc.add_paragraph()  # spacing after table
                continue

        # --- Headings ---
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # --- Blockquote ---
        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
            i += 1
            continue

        # --- Checkbox list item ---
        m = re.match(r"^(\s*)-\s*\[(.)\]\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            checked = m.group(2).strip().lower() == "x"
            text = m.group(3)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 + indent_spaces * 0.15)
            # Box character
            box = p.add_run("☒ " if checked else "☐ ")
            box.font.size = Pt(11)
            if checked:
                box.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
            add_inline_runs(p, text)
            i += 1
            continue

        # --- Bullet list item ---
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + indent_spaces * 0.15)
            add_inline_runs(p, text)
            i += 1
            continue

        # --- Numbered list item ---
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Inches(0.25 + indent_spaces * 0.15)
            add_inline_runs(p, text)
            i += 1
            continue

        # --- Blank line ---
        if not stripped:
            i += 1
            continue

        # --- Plain paragraph ---
        # Collect consecutive non-blank, non-special lines into one paragraph
        para_lines: list[str] = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_s = nxt.strip()
            if (not nxt_s
                    or nxt_s.startswith(("#", "-", "*", ">", "|", "```", "---"))
                    or re.match(r"^\s*\d+\.\s", nxt)):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_inline_runs(p, para_text)

    # ----- Save -----
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ---------------------------------------------------------------------------
# Optional: MD -> PDF via reportlab (simpler than DOCX renderer)
# ---------------------------------------------------------------------------

def md_to_pdf(md_text: str, out_path: Path) -> None:
    """Render the markdown into a paginated PDF.

    Uses reportlab Platypus.  Same subset as md_to_docx renders identically
    where reasonable.  PDF is a secondary output; DOCX is the primary.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (Paragraph, SimpleDocTemplate, Spacer,
                                    Table, TableStyle, Preformatted, HRFlowable)

    styles = getSampleStyleSheet()
    body = ParagraphStyle("body", parent=styles["BodyText"],
                          fontName="Helvetica", fontSize=9, leading=12,
                          spaceAfter=4)
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=16, leading=20,
                        spaceBefore=14, spaceAfter=8, textColor=colors.HexColor("#1A4F8B"))
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=13, leading=16,
                        spaceBefore=10, spaceAfter=6, textColor=colors.HexColor("#2C5F9B"))
    h3 = ParagraphStyle("h3", parent=styles["Heading3"], fontSize=11, leading=14,
                        spaceBefore=8, spaceAfter=4, textColor=colors.HexColor("#3E6FAB"))
    code_style = ParagraphStyle("code", parent=body, fontName="Courier",
                                fontSize=8, leading=10,
                                backColor=colors.HexColor("#F4F4F4"),
                                borderColor=colors.HexColor("#CCCCCC"),
                                borderWidth=0.5, borderPadding=4,
                                spaceBefore=4, spaceAfter=4)
    quote_style = ParagraphStyle("quote", parent=body, fontName="Helvetica-Oblique",
                                 textColor=colors.HexColor("#505050"), leftIndent=18)

    def md_inline_to_rl(text: str) -> str:
        """Convert **bold** / `code` to ReportLab paragraph markup."""
        out = re.sub(r"\*\*([^*]+)\*\*", r"<b>\1</b>", text)
        out = re.sub(r"`([^`]+)`",
                     r'<font face="Courier" size="8">\1</font>', out)
        # Checkbox glyphs survive
        return out

    flowables: list = []
    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped == "---":
            flowables.append(HRFlowable(width="100%", thickness=0.5,
                                        color=colors.HexColor("#999999"),
                                        spaceBefore=6, spaceAfter=6))
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            flowables.append(Preformatted("\n".join(code_lines), code_style))
            continue

        if ("|" in stripped and i + 1 < n
                and re.match(r"^\s*\|?\s*[-:|\s]+\|", lines[i + 1])):
            rows: list[list[str]] = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append([c.strip() for c in
                             lines[i].strip().strip("|").split("|")])
                i += 1
            if len(rows) >= 2:
                header = rows[0]
                body_rows = rows[2:]
                data = [[Paragraph(md_inline_to_rl(c) or "&nbsp;", body)
                         for c in header]]
                for r in body_rows:
                    data.append([Paragraph(md_inline_to_rl(c) or "&nbsp;",
                                           body) for c in r])
                tbl = Table(data, repeatRows=1, colWidths=None,
                            hAlign="LEFT")
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0),
                     colors.HexColor("#D9E1F2")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.25,
                     colors.HexColor("#999999")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                flowables.append(tbl)
                flowables.append(Spacer(1, 6))
                continue

        if stripped.startswith("### "):
            flowables.append(Paragraph(md_inline_to_rl(stripped[4:]), h3))
            i += 1
            continue
        if stripped.startswith("## "):
            flowables.append(Paragraph(md_inline_to_rl(stripped[3:]), h2))
            i += 1
            continue
        if stripped.startswith("# "):
            flowables.append(Paragraph(md_inline_to_rl(stripped[2:]), h1))
            i += 1
            continue

        if stripped.startswith("> "):
            flowables.append(Paragraph(md_inline_to_rl(stripped[2:]),
                                       quote_style))
            i += 1
            continue

        m = re.match(r"^(\s*)-\s*\[(.)\]\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            checked = m.group(2).strip().lower() == "x"
            text = m.group(3)
            box = "☒" if checked else "☐"
            color = '<font color="#008000">' if checked else ""
            close_color = "</font>" if checked else ""
            indent_em = 0.18 + indent_spaces * 0.12
            para_style = ParagraphStyle(
                f"cb{indent_spaces}", parent=body,
                leftIndent=indent_em * inch, firstLineIndent=-0.18 * inch)
            flowables.append(Paragraph(
                f"{color}{box}{close_color} {md_inline_to_rl(text)}",
                para_style))
            i += 1
            continue

        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            text = m.group(2)
            indent_em = 0.18 + indent_spaces * 0.12
            para_style = ParagraphStyle(
                f"bul{indent_spaces}", parent=body,
                leftIndent=indent_em * inch, firstLineIndent=-0.18 * inch)
            flowables.append(Paragraph(f"• {md_inline_to_rl(text)}",
                                       para_style))
            i += 1
            continue

        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            indent_spaces = len(m.group(1))
            num = m.group(2)
            text = m.group(3)
            indent_em = 0.18 + indent_spaces * 0.12
            para_style = ParagraphStyle(
                f"num{indent_spaces}", parent=body,
                leftIndent=indent_em * inch, firstLineIndent=-0.18 * inch)
            flowables.append(Paragraph(f"{num}. {md_inline_to_rl(text)}",
                                       para_style))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        # Plain paragraph (collect consecutive lines)
        para_lines: list[str] = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            nxt_s = nxt.strip()
            if (not nxt_s
                    or nxt_s.startswith(("#", "-", "*", ">", "|", "```",
                                         "---"))
                    or re.match(r"^\s*\d+\.\s", nxt)):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        flowables.append(Paragraph(md_inline_to_rl(para_text), body))

    out_path.parent.mkdir(parents=True, exist_ok=True)

    def _footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#808080"))
        canvas.drawString(
            0.5 * inch, 0.3 * inch,
            f"Lyra-cpp Execution Plan — auto-generated "
            f"{_dt.datetime.now().strftime('%Y-%m-%d %H:%M')} "
            f"— page {doc_.page}")
        canvas.restoreState()

    pdf_doc = SimpleDocTemplate(
        str(out_path), pagesize=LETTER,
        leftMargin=0.5 * inch, rightMargin=0.5 * inch,
        topMargin=0.5 * inch, bottomMargin=0.5 * inch,
        title="Lyra-cpp TX Execution Plan",
        author="N8SDR + Claude")
    pdf_doc.build(flowables, onFirstPage=_footer, onLaterPages=_footer)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Sync EXECUTION_PLAN.md to DOCX (+ optional PDF).")
    parser.add_argument("--pdf", action="store_true",
                        help="Also generate EXECUTION_PLAN.pdf")
    parser.add_argument("--md", type=Path, default=MD_PATH,
                        help=f"Source markdown (default: {MD_PATH})")
    parser.add_argument("--docx", type=Path, default=DOCX_PATH,
                        help=f"Destination DOCX (default: {DOCX_PATH})")
    parser.add_argument("--pdf-path", type=Path, default=PDF_PATH,
                        help=f"Destination PDF (default: {PDF_PATH})")
    args = parser.parse_args(argv)

    if not args.md.exists():
        print(f"ERROR: source markdown not found: {args.md}",
              file=sys.stderr)
        return 2

    md_text = args.md.read_text(encoding="utf-8")

    print(f"[sync] reading {args.md}")
    print(f"[sync] writing DOCX -> {args.docx}")
    md_to_docx(md_text, args.docx)
    print(f"[sync] DOCX OK ({args.docx.stat().st_size:,} bytes)")

    if args.pdf:
        print(f"[sync] writing PDF -> {args.pdf_path}")
        md_to_pdf(md_text, args.pdf_path)
        print(f"[sync] PDF OK ({args.pdf_path.stat().st_size:,} bytes)")

    return 0


if __name__ == "__main__":
    sys.exit(main())
